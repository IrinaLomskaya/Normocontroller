from __future__ import annotations

import json
import re
import uuid
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Iterator, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.section import Section
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


# ============================================================
# Helpers
# ============================================================

PLACEHOLDER_PATTERNS = [
    r"_{3,}",
    r"<[^<>]{0,100}>",
    r"\[[^\[\]]{0,100}\]",
    r"\{\{[^{}]{0,100}\}\}",
    r"\bуказать\b",
    r"\bзаполнить\b",
]


def gen_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


def normalize_text(text: str) -> str:
    if text is None:
        return ""
    text = text.replace("\xa0", " ")
    text = text.replace("\r", "\n")
    text = text.replace("—", "-")
    text = text.replace("–", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def emu_to_pt(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value.pt)
    except Exception:
        try:
            return float(value) / 12700.0
        except Exception:
            return None


def safe_enum_name(value: Any) -> Optional[str]:
    if value is None:
        return None
    try:
        return value.name
    except Exception:
        return str(value)


def detect_placeholders(text: str) -> list[dict[str, Any]]:
    findings = []
    if not text:
        return findings

    for pattern in PLACEHOLDER_PATTERNS:
        for m in re.finditer(pattern, text, flags=re.IGNORECASE):
            findings.append(
                {
                    "pattern": pattern,
                    "match": m.group(0),
                    "start": m.start(),
                    "end": m.end(),
                }
            )
    return findings


def guess_heading_level(style_name: Optional[str], text: str) -> Optional[int]:
    if style_name:
        s = style_name.lower().strip()
        m = re.search(r"heading\s*([1-9])", s)
        if m:
            return int(m.group(1))
        m = re.search(r"заголовок\s*([1-9])", s)
        if m:
            return int(m.group(1))

    text = text.strip()
    if re.match(r"^\d+\.\s+\S+", text):
        return 1
    if re.match(r"^\d+\.\d+\.\s+\S+", text):
        return 2
    if re.match(r"^\d+\.\d+\.\d+\.\s+\S+", text):
        return 3
    return None


def is_heading(style_name: Optional[str], text: str) -> tuple[bool, Optional[int]]:
    level = guess_heading_level(style_name, text)
    if level is not None:
        return True, level

    if style_name:
        s = style_name.lower()
        if "heading" in s or "заголовок" in s:
            return True, level

    return False, None


def paragraph_numbering_info(paragraph: Paragraph) -> Optional[dict[str, Any]]:
    p = paragraph._p
    pPr = p.pPr
    if pPr is None or pPr.numPr is None:
        return None

    num_pr = pPr.numPr
    ilvl = None
    num_id = None

    if num_pr.ilvl is not None:
        try:
            ilvl = int(num_pr.ilvl.val)
        except Exception:
            ilvl = None

    if num_pr.numId is not None:
        try:
            num_id = int(num_pr.numId.val)
        except Exception:
            num_id = None

    return {"num_id": num_id, "ilvl": ilvl}


def iter_block_items(parent: DocxDocument) -> Iterator[Paragraph | Table]:
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def iter_cell_block_items(cell: _Cell) -> Iterator[Paragraph | Table]:
    for child in cell._tc.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, cell)
        elif isinstance(child, CT_Tbl):
            yield Table(child, cell)


def read_xml_attr(element: Any, xpath: str, attr_name: str, nsmap: dict[str, str]) -> Optional[str]:
    try:
        nodes = element.xpath(xpath, namespaces=nsmap)
        if not nodes:
            return None
        node = nodes[0]
        return node.get(attr_name)
    except Exception:
        return None


def safe_get_row_cells(row) -> list:
    """
    Безопасно получить ячейки строки таблицы.
    python-docx может падать на row.cells для merge/vMerge таблиц.
    """
    try:
        return list(row.cells)
    except Exception:
        return []


# ============================================================
# Extractors
# ============================================================

def extract_style_info(obj: Any) -> dict[str, Any]:
    style = getattr(obj, "style", None)
    if style is None:
        return {
            "name": None,
            "style_id": None,
            "type": None,
            "base_style": None,
        }

    base_style = getattr(style, "base_style", None)
    return {
        "name": getattr(style, "name", None),
        "style_id": getattr(style, "style_id", None),
        "type": str(getattr(style, "type", None)) if getattr(style, "type", None) is not None else None,
        "base_style": getattr(base_style, "name", None) if base_style is not None else None,
    }


def extract_run_info(paragraph: Paragraph) -> list[dict[str, Any]]:
    runs = []

    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        if not text:
            continue

        font = run.font
        runs.append(
            {
                "run_index": idx,
                "text": text,
                "text_normalized": normalize_text(text),
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "all_caps": getattr(font, "all_caps", None),
                "small_caps": getattr(font, "small_caps", None),
                "subscript": getattr(font, "subscript", None),
                "superscript": getattr(font, "superscript", None),
                "font_name": getattr(font, "name", None),
                "font_size_pt": emu_to_pt(getattr(font, "size", None)),
                "highlight_color": str(getattr(font, "highlight_color", None))
                if getattr(font, "highlight_color", None) is not None else None,
                "color_rgb": str(font.color.rgb) if getattr(font.color, "rgb", None) is not None else None,
            }
        )

    return runs


def extract_paragraph_format(paragraph: Paragraph) -> dict[str, Any]:
    pf = paragraph.paragraph_format

    return {
        "alignment": safe_enum_name(paragraph.alignment),
        "left_indent_pt": emu_to_pt(pf.left_indent),
        "right_indent_pt": emu_to_pt(pf.right_indent),
        "first_line_indent_pt": emu_to_pt(pf.first_line_indent),
        "space_before_pt": emu_to_pt(pf.space_before),
        "space_after_pt": emu_to_pt(pf.space_after),
        "line_spacing": pf.line_spacing,
        "line_spacing_rule": safe_enum_name(getattr(pf, "line_spacing_rule", None)),
        "keep_together": getattr(pf, "keep_together", None),
        "keep_with_next": getattr(pf, "keep_with_next", None),
        "page_break_before": getattr(pf, "page_break_before", None),
        "widow_control": getattr(pf, "widow_control", None),
    }


def extract_section_info(section: Section, section_index: int) -> dict[str, Any]:
    return {
        "section_index": section_index,
        "start_type": safe_enum_name(getattr(section, "start_type", None)),
        "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
        "page_width_pt": emu_to_pt(section.page_width),
        "page_height_pt": emu_to_pt(section.page_height),
        "margin_left_pt": emu_to_pt(section.left_margin),
        "margin_right_pt": emu_to_pt(section.right_margin),
        "margin_top_pt": emu_to_pt(section.top_margin),
        "margin_bottom_pt": emu_to_pt(section.bottom_margin),
        "header_distance_pt": emu_to_pt(section.header_distance),
        "footer_distance_pt": emu_to_pt(section.footer_distance),
        "gutter_pt": emu_to_pt(getattr(section, "gutter", None)),
    }


def extract_table_format(table: Table) -> dict[str, Any]:
    try:
        style_name = getattr(table.style, "name", None) if getattr(table, "style", None) is not None else None
    except Exception:
        style_name = None

    try:
        alignment = safe_enum_name(table.alignment)
    except Exception:
        alignment = None

    try:
        autofit = table.autofit
    except Exception:
        autofit = None

    return {
        "style": style_name,
        "alignment": alignment,
        "autofit": autofit,
    }


def extract_cell_format(cell: _Cell) -> dict[str, Any]:
    tcPr = cell._tc.tcPr
    grid_span = None
    v_merge = None
    width_pt = None
    fill = None
    vertical_alignment = safe_enum_name(getattr(cell, "vertical_alignment", None))

    if tcPr is not None:
        if tcPr.gridSpan is not None:
            try:
                grid_span = int(tcPr.gridSpan.val)
            except Exception:
                grid_span = None

        if tcPr.vMerge is not None:
            try:
                v_merge = str(tcPr.vMerge.val)
            except Exception:
                v_merge = "continue"

        if tcPr.tcW is not None:
            try:
                width_pt = float(tcPr.tcW.w) / 20.0
            except Exception:
                width_pt = None

        try:
            shd = tcPr.shd
            if shd is not None:
                fill = shd.fill
        except Exception:
            fill = None

    return {
        "vertical_alignment": vertical_alignment,
        "grid_span": grid_span,
        "v_merge": v_merge,
        "width_pt": width_pt,
        "fill": fill,
    }


def extract_headers_footers(section: Section, section_index: int) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    headers = []
    footers = []

    try:
        header_paragraphs = []
        for p in section.header.paragraphs:
            t = normalize_text(p.text or "")
            if t:
                header_paragraphs.append(t)

        headers.append(
            {
                "section_index": section_index,
                "kind": "default",
                "text": "\n".join(header_paragraphs),
                "paragraphs": header_paragraphs,
            }
        )
    except Exception:
        pass

    try:
        footer_paragraphs = []
        for p in section.footer.paragraphs:
            t = normalize_text(p.text or "")
            if t:
                footer_paragraphs.append(t)

        footers.append(
            {
                "section_index": section_index,
                "kind": "default",
                "text": "\n".join(footer_paragraphs),
                "paragraphs": footer_paragraphs,
            }
        )
    except Exception:
        pass

    return headers, footers


# ============================================================
# Dataclasses
# ============================================================

@dataclass
class Block:
    block_id: str
    order_index: int
    block_type: str
    text_raw: str
    text_normalized: str
    style: dict[str, Any]
    level: Optional[int]
    section_path: list[str]
    parent_block_id: Optional[str]
    numbering: Optional[dict[str, Any]]
    paragraph_format: Optional[dict[str, Any]]
    runs: Optional[list[dict[str, Any]]]
    formatting_summary: Optional[dict[str, Any]]
    layout_context: dict[str, Any]
    placeholders: list[dict[str, Any]]
    source_ref: dict[str, Any]


@dataclass
class TableCellData:
    row_idx: int
    col_idx: int
    text_raw: str
    text_normalized: str
    cell_format: dict[str, Any]
    paragraphs: list[dict[str, Any]]


@dataclass
class ParsedTable:
    table_id: str
    block_id: str
    rows_count: int
    cols_count: int
    section_path: list[str]
    table_format: dict[str, Any]
    cells: list[dict[str, Any]]
    schema_signature: dict[str, Any]


# ============================================================
# Main parser
# ============================================================

class DocxNormalizer:
    def __init__(self) -> None:
        self.section_stack: list[tuple[int, str]] = []

    def _update_section_path(self, heading_text: str, level: int) -> list[str]:
        while self.section_stack and self.section_stack[-1][0] >= level:
            self.section_stack.pop()
        self.section_stack.append((level, heading_text))
        return [title for _, title in self.section_stack]

    def _current_section_path(self) -> list[str]:
        return [title for _, title in self.section_stack]

    def _resolve_section_idx_for_block(self, doc: DocxDocument, order_index: int) -> int:
        return 0 if len(doc.sections) > 0 else -1

    def _layout_context_from_section(self, sections_data: list[dict[str, Any]], section_idx: int) -> dict[str, Any]:
        if 0 <= section_idx < len(sections_data):
            return sections_data[section_idx]
        return {"section_index": section_idx}

    def _extract_formatting_summary(self, runs: list[dict[str, Any]]) -> dict[str, Any]:
        return {
            "has_bold": any(r.get("bold") for r in runs),
            "has_italic": any(r.get("italic") for r in runs),
            "has_underline": any(r.get("underline") for r in runs),
            "font_names": sorted(list({r["font_name"] for r in runs if r.get("font_name")})),
            "font_sizes_pt": sorted(list({r["font_size_pt"] for r in runs if r.get("font_size_pt") is not None})),
        }

    def _parse_table_cells(self, table: Table) -> tuple[list[dict[str, Any]], dict[str, Any], int, int]:
        cells_out: list[dict[str, Any]] = []
        header_row: list[str] = []

        safe_rows: list[list[_Cell]] = []
        for row in table.rows:
            row_cells = safe_get_row_cells(row)
            safe_rows.append(row_cells)

        rows_count = len(safe_rows)
        cols_count = max((len(r) for r in safe_rows), default=0)

        for row_idx, row_cells in enumerate(safe_rows):
            for col_idx, cell in enumerate(row_cells):
                try:
                    text_raw = cell.text or ""
                except Exception:
                    text_raw = ""

                text_normalized = normalize_text(text_raw)

                paragraphs_data = []
                try:
                    cell_paragraphs = list(cell.paragraphs)
                except Exception:
                    cell_paragraphs = []

                for p in cell_paragraphs:
                    try:
                        p_text = p.text or ""
                    except Exception:
                        p_text = ""

                    p_norm = normalize_text(p_text)

                    try:
                        style_info = extract_style_info(p)
                    except Exception:
                        style_info = {
                            "name": None,
                            "style_id": None,
                            "type": None,
                            "base_style": None,
                        }

                    try:
                        paragraph_format = extract_paragraph_format(p)
                    except Exception:
                        paragraph_format = {}

                    try:
                        runs = extract_run_info(p)
                    except Exception:
                        runs = []

                    paragraphs_data.append(
                        {
                            "text_raw": p_text,
                            "text_normalized": p_norm,
                            "style": style_info,
                            "paragraph_format": paragraph_format,
                            "runs": runs,
                        }
                    )

                try:
                    cell_format = extract_cell_format(cell)
                except Exception:
                    cell_format = {
                        "vertical_alignment": None,
                        "grid_span": None,
                        "v_merge": None,
                        "width_pt": None,
                        "fill": None,
                    }

                cell_data = TableCellData(
                    row_idx=row_idx,
                    col_idx=col_idx,
                    text_raw=text_raw,
                    text_normalized=text_normalized,
                    cell_format=cell_format,
                    paragraphs=paragraphs_data,
                )
                cells_out.append(asdict(cell_data))

                if row_idx == 0:
                    header_row.append(text_normalized)

        schema_signature = {
            "rows_count": rows_count,
            "cols_count": cols_count,
            "header": header_row,
        }

        return cells_out, schema_signature, rows_count, cols_count

    def parse(self, file_path: str | Path) -> dict[str, Any]:
        file_path = Path(file_path)
        doc = Document(str(file_path))

        sections_data = [
            extract_section_info(section, idx)
            for idx, section in enumerate(doc.sections)
        ]

        headers = []
        footers = []
        for idx, section in enumerate(doc.sections):
            hs, fs = extract_headers_footers(section, idx)
            headers.extend(hs)
            footers.extend(fs)

        blocks: list[Block] = []
        tables: list[ParsedTable] = []

        order_index = 0

        for item in iter_block_items(doc):
            section_idx = self._resolve_section_idx_for_block(doc, order_index)
            layout_context = self._layout_context_from_section(sections_data, section_idx)

            if isinstance(item, Paragraph):
                text_raw = item.text or ""
                text_normalized = normalize_text(text_raw)

                if not text_normalized:
                    continue

                style_info = extract_style_info(item)
                heading_flag, heading_level = is_heading(style_info["name"], text_normalized)

                if heading_flag:
                    section_path = self._update_section_path(text_normalized, heading_level or 1)
                    block_type = "heading"
                    level = heading_level
                else:
                    section_path = self._current_section_path()
                    block_type = "paragraph"
                    level = None

                runs = extract_run_info(item)
                paragraph_format = extract_paragraph_format(item)

                block = Block(
                    block_id=gen_id("blk"),
                    order_index=order_index,
                    block_type=block_type,
                    text_raw=text_raw,
                    text_normalized=text_normalized,
                    style=style_info,
                    level=level,
                    section_path=section_path,
                    parent_block_id=None,
                    numbering=paragraph_numbering_info(item),
                    paragraph_format=paragraph_format,
                    runs=runs,
                    formatting_summary=self._extract_formatting_summary(runs),
                    layout_context=layout_context,
                    placeholders=detect_placeholders(text_raw),
                    source_ref={
                        "part": "body",
                        "kind": "paragraph",
                        "order_index": order_index,
                    },
                )
                blocks.append(block)
                order_index += 1

            elif isinstance(item, Table):
                section_path = self._current_section_path()

                try:
                    cells_out, schema_signature, rows_count, cols_count = self._parse_table_cells(item)

                    table_matrix: dict[int, dict[int, str]] = {}
                    for cell_info in cells_out:
                        r = cell_info["row_idx"]
                        c = cell_info["col_idx"]
                        txt = cell_info["text_normalized"] or ""

                        if r not in table_matrix:
                            table_matrix[r] = {}
                        table_matrix[r][c] = txt

                    table_text_parts = []
                    for r in range(rows_count):
                        row_cells = []
                        for c in range(cols_count):
                            row_cells.append(table_matrix.get(r, {}).get(c, ""))
                        row_text = " | ".join(row_cells).strip()
                        table_text_parts.append(row_text)

                    table_text = "\n".join(table_text_parts).strip()

                    block_id = gen_id("blk")

                    block = Block(
                        block_id=block_id,
                        order_index=order_index,
                        block_type="table",
                        text_raw=table_text,
                        text_normalized=normalize_text(table_text),
                        style=extract_style_info(item),
                        level=None,
                        section_path=section_path,
                        parent_block_id=None,
                        numbering=None,
                        paragraph_format=None,
                        runs=None,
                        formatting_summary=None,
                        layout_context=layout_context,
                        placeholders=detect_placeholders(table_text),
                        source_ref={
                            "part": "body",
                            "kind": "table",
                            "order_index": order_index,
                        },
                    )
                    blocks.append(block)

                    parsed_table = ParsedTable(
                        table_id=gen_id("tbl"),
                        block_id=block_id,
                        rows_count=rows_count,
                        cols_count=cols_count,
                        section_path=section_path,
                        table_format=extract_table_format(item),
                        cells=cells_out,
                        schema_signature=schema_signature,
                    )
                    tables.append(parsed_table)

                    order_index += 1

                except Exception as e:
                    block_id = gen_id("blk")

                    block = Block(
                        block_id=block_id,
                        order_index=order_index,
                        block_type="table",
                        text_raw="",
                        text_normalized="",
                        style=extract_style_info(item),
                        level=None,
                        section_path=section_path,
                        parent_block_id=None,
                        numbering=None,
                        paragraph_format=None,
                        runs=None,
                        formatting_summary=None,
                        layout_context=layout_context,
                        placeholders=[],
                        source_ref={
                            "part": "body",
                            "kind": "table",
                            "order_index": order_index,
                            "parse_error": str(e),
                        },
                    )
                    blocks.append(block)

                    parsed_table = ParsedTable(
                        table_id=gen_id("tbl"),
                        block_id=block_id,
                        rows_count=0,
                        cols_count=0,
                        section_path=section_path,
                        table_format={
                            "style": getattr(item.style, "name", None) if getattr(item, "style", None) is not None else None,
                            "alignment": None,
                            "autofit": None,
                            "parse_error": str(e),
                        },
                        cells=[],
                        schema_signature={
                            "rows_count": 0,
                            "cols_count": 0,
                            "header": [],
                            "parse_error": str(e),
                        },
                    )
                    tables.append(parsed_table)

                    order_index += 1

        core_properties = {
            "author": getattr(doc.core_properties, "author", None),
            "title": getattr(doc.core_properties, "title", None),
            "subject": getattr(doc.core_properties, "subject", None),
            "category": getattr(doc.core_properties, "category", None),
            "comments": getattr(doc.core_properties, "comments", None),
            "created": str(getattr(doc.core_properties, "created", None))
            if getattr(doc.core_properties, "created", None) else None,
            "modified": str(getattr(doc.core_properties, "modified", None))
            if getattr(doc.core_properties, "modified", None) else None,
            "last_modified_by": getattr(doc.core_properties, "last_modified_by", None),
        }

        result = {
            "document": {
                "document_id": gen_id("doc"),
                "filename": file_path.name,
                "source_path": str(file_path),
                "core_properties": core_properties,
                "stats": {
                    "sections_count": len(sections_data),
                    "blocks_count": len(blocks),
                    "tables_count": len(tables),
                    "headers_count": len(headers),
                    "footers_count": len(footers),
                },
            },
            "sections": sections_data,
            "blocks": [asdict(b) for b in blocks],
            "tables": [asdict(t) for t in tables],
            "headers": headers,
            "footers": footers,
        }
        return result


# ============================================================
# CLI usage
# ============================================================

# if __name__ == "__main__":
#     parser = DocxNormalizer()
#     parsed = parser.parse("sample.docx")

#     with open("parsed_sample_expanded2.json", "w", encoding="utf-8") as f:
#         json.dump(parsed, f, ensure_ascii=False, indent=2)

#     print("Saved to parsed_sample_expanded2.json")